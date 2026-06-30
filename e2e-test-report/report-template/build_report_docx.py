# -*- coding: utf-8 -*-
"""Beautiful QA report docx builder (python-docx) — SELF-CONTAINED.
Excellence criteria: cover + score badge, color-coded rubric table, feature inventory,
detail test per criterion (card C1-C7), sized captioned screenshots, professional typography, page numbers.

Đọc dữ liệu từ sample_data.py (cùng folder). Thay data sản phẩm thật của bạn vào sample_data.py,
rồi chạy:

    python3 build_report_docx.py

Output: <QA_REPORT_DIR>/QA-Report-<ProductName>.docx  (default ./test_report_output)
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from sample_data import PRODUCT, TD, INVENTORY, CRIT, W
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

B = os.path.expanduser(os.environ.get("QA_SCREENSHOT_DIR", os.path.join(os.path.dirname(__file__), "samples")))
RD = os.path.expanduser(os.environ.get("QA_REPORT_DIR", "./test_report_output"))

# Palette
NAVY = RGBColor(0x1F, 0x38, 0x64); TEAL = RGBColor(0x0F, 0x6B, 0x7B)
GREEN = RGBColor(0x1E, 0x7A, 0x32); AMBER = RGBColor(0xB8, 0x6A, 0x00); RED = RGBColor(0xB3, 0x1B, 0x1B)
GRAY = RGBColor(0x59, 0x59, 0x59); LIGHT = RGBColor(0x6B, 0x72, 0x80); WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x22, 0x22, 0x22)
NAVY_HEX = "1F3864"; TEAL_HEX = "0F6B7B"; LIGHTBG_HEX = "EEF2F7"; ROWBG_HEX = "F7F9FC"


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr(); m = OxmlElement('w:tcMar')
    for tag, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)


def cell_text(cell, text, bold=False, color=None, size=10, align='left', italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    for i, line in enumerate(str(text).split("\n")):
        if i > 0: p.add_run().add_break()
        r = p.add_run(line); r.bold = bold; r.italic = italic; r.font.size = Pt(size); r.font.name = 'Calibri'
        if color is not None: r.font.color.rgb = color
    return cell


def score_color(v):
    if v is None: return GRAY
    return GREEN if v >= 70 else (AMBER if v >= 50 else RED)


def score_hex(v):
    if v is None: return "8A8F98"
    return "1E7A32" if v >= 70 else ("B86A00" if v >= 50 else "B31B1B")


def add_page_number_footer(section, product_name):
    footer = section.footer; p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"QA Report · {product_name} · e2e-test-report · trang ")
    r.font.size = Pt(8); r.font.color.rgb = LIGHT; r.font.name = 'Calibri'
    fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
    run = OxmlElement('w:r'); rpr = OxmlElement('w:rPr'); sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '16'); rpr.append(sz)
    run.append(rpr); t = OxmlElement('w:t'); t.text = "1"; run.append(t); fld1.append(run); p._p.append(fld1)


def add_hr(doc, color_hex=TEAL_HEX, size=12):
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), color_hex); pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.space_before = Pt(2)


def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    add_hr(doc, TEAL_HEX, 8)
    return p


def add_image(doc, path, caption, width_in=6.0):
    if not (path and os.path.exists(path)):
        return
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(caption); cr.italic = True; cr.font.size = Pt(8.5); cr.font.color.rgb = LIGHT; cr.font.name = 'Calibri'
    c.paragraph_format.space_after = Pt(8)


def set_base_style(doc):
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5); st.font.color.rgb = DARK
    for sec in doc.sections:
        sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7); sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)


def aggregate(levels, no_cb=False):
    w = dict(W)
    if no_cb:
        w["C3"] = w["C3"] + 10; w["C5"] = w["C5"] + 5; w["C4"] = 0
    total = 0; contrib = {}; used_w = 0
    for c, lvl in levels.items():
        wt = w[c]
        if lvl == 0:
            contrib[c] = (0, wt, 0, 0); continue
        sub = (lvl / 5) * 100; d = sub * wt / 100; total += d; used_w += wt
        contrib[c] = (round(sub), wt, lvl, round(d, 1))
    final = round(total * 100 / used_w, 1) if used_w else 0
    return final, contrib, used_w, w


def build_cover(doc, t, final, cov, conf, band):
    # top color band
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = True
    cell = tbl.rows[0].cells[0]; shade(cell, NAVY_HEX); set_cell_margins(cell, 220, 220, 200, 200)
    cell_text(cell, "BÁO CÁO QA · SẢN PHẨM PHẦN MỀM", bold=True, color=WHITE, size=11)
    sub = cell.paragraphs[0]
    sub2 = sub.add_run(f"\nE2E Test Report · {t['date']}")
    sub2.font.color.rgb = RGBColor(0xBD, 0xD4, 0xEA); sub2.font.size = Pt(9); sub2.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    # product name big
    p = doc.add_paragraph(); r = p.add_run(t['name']); r.bold = True; r.font.size = Pt(26)
    r.font.color.rgb = NAVY; r.font.name = 'Calibri'; p.paragraph_format.space_after = Pt(2)
    p2 = doc.add_paragraph(); r2 = p2.add_run(t['topic']); r2.font.size = Pt(11); r2.italic = True
    r2.font.color.rgb = TEAL; r2.font.name = 'Calibri'; p2.paragraph_format.space_after = Pt(10)
    # info table
    info = doc.add_table(rows=4, cols=2); info.allow_autofit = True
    rows = [("URL kiểm thử", t['url']),
            ("Phương pháp", "Feature Explore theo role → Bảng tính năng → test từng feature → chấm rubric 7 hạng mục (mỗi điểm = test + ảnh chụp)"),
            ("Chế độ", "Report-only (không sửa code)"),
            ("Ngày kiểm thử", t['date'])]
    for i, (k, v) in enumerate(rows):
        c0 = info.rows[i].cells[0]; c1 = info.rows[i].cells[1]
        shade(c0, LIGHTBG_HEX); set_cell_margins(c0); set_cell_margins(c1)
        cell_text(c0, k, bold=True, color=NAVY, size=9.5); cell_text(c1, v, size=9.5)
        c0.width = Inches(1.4); c1.width = Inches(5.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    # Score badge box
    if final is not None:
        bt = doc.add_table(rows=1, cols=3); bt.alignment = WD_TABLE_ALIGNMENT.CENTER
        cells = bt.rows[0].cells
        for c in cells: set_cell_margins(c, 160, 160, 120, 120)
        sc = cells[0]; shade(sc, score_hex(final)); cell_text(sc, "", size=10)
        sc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = sc.paragraphs[0].add_run(f"{final}"); rr.bold = True; rr.font.size = Pt(30)
        rr.font.color.rgb = WHITE; rr.font.name = 'Calibri'
        rr2 = sc.paragraphs[0].add_run(" /100"); rr2.font.size = Pt(12); rr2.font.color.rgb = WHITE; rr2.font.name = 'Calibri'
        sc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl = sc.add_paragraph(); lr = lbl.add_run("ĐIỂM TỔNG"); lr.font.size = Pt(8); lr.font.color.rgb = WHITE; lr.bold = True
        lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cv = cells[1]; shade(cv, ROWBG_HEX); cv.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cv.paragraphs[0].add_run(f"{cov}%"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
        lb = cv.add_paragraph(); lr = lb.add_run("ĐỘ PHỦ"); lr.font.size = Pt(8); lr.font.color.rgb = GRAY; lr.bold = True
        lb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lb2 = cv.add_paragraph(); lr2 = lb2.add_run(f"{conf} tin cậy"); lr2.font.size = Pt(8); lr2.font.color.rgb = LIGHT
        lb2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bd = cells[2]; shade(bd, score_hex(final)); bd.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rb = bd.paragraphs[0].add_run(band.split("—")[0].strip()); rb.bold = True; rb.font.size = Pt(12)
        rb.font.color.rgb = WHITE; rb.font.name = 'Calibri'
        lb = bd.add_paragraph(); lr = lb.add_run(band.split("—")[-1].strip() if "—" in band else band)
        lr.font.size = Pt(8); lr.font.color.rgb = WHITE; lb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def rubric_table(doc, levels, no_cb, t):
    final, contrib, used_w, w = aggregate(levels, no_cb)
    tbl = doc.add_table(rows=1, cols=6); tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    heads = ["Hạng mục", "Trọng số", "Cấp", "Điểm", "Bằng chứng kiểm thử", "Ảnh"]
    for i, htext in enumerate(heads):
        shade(hdr[i], NAVY_HEX); set_cell_margins(hdr[i])
        cell_text(hdr[i], htext, bold=True, color=WHITE, size=9, align='center' if i in (1, 2, 3) else 'left')
    for c in ["C1", "C2", "C3", "C4", "C5", "C6", "C7"]:
        sub, wt, lvl, d = contrib[c]; row = tbl.add_row().cells
        for cc in row: set_cell_margins(cc)
        cell_text(row[0], CRIT[c], bold=True, size=9, color=NAVY)
        wtshow = f"{w[c]}%" if not (no_cb and c == "C4") else "N/A"
        cell_text(row[1], wtshow, size=9, align='center', color=GRAY)
        if lvl == 0:
            cell_text(row[2], "—", size=9, align='center', color=LIGHT); shade(row[2], "EDEDED")
            cell_text(row[3], "UNV", size=8.5, align='center', color=LIGHT); shade(row[3], "EDEDED")
        else:
            cell_text(row[2], f"{lvl}/5", size=9, align='center', bold=True); shade(row[2], score_hex(sub))
            row[2].paragraphs[0].runs[0].font.color.rgb = WHITE
            cell_text(row[3], f"{sub}", size=9, align='center')
        cell_text(row[4], t["rat"].get(c, "")[:120], size=8.5, color=RGBColor(0x33, 0x33, 0x33))
        cell_text(row[5], t["shot"].get(c) or "—", size=8, color=GRAY, align='center')
    tr = tbl.add_row().cells
    for cc in tr: shade(cc, LIGHTBG_HEX); set_cell_margins(cc)
    cell_text(tr[0], "TỔNG", bold=True, color=NAVY, size=10)
    cell_text(tr[1], f"{used_w}%", bold=True, size=9, align='center', color=NAVY)
    cell_text(tr[2], "", size=9)
    cell_text(tr[3], f"{final}", bold=True, size=11, align='center', color=score_color(final))
    unver = sum(1 for l in levels.values() if l == 0)
    conf = 0.55 if unver >= 4 else (0.75 if unver >= 2 else 0.92)
    cell_text(tr[4], f"confidence {conf}", size=8.5, italic=True, color=GRAY)
    cell_text(tr[5], "", size=8)
    for r_ in tbl.rows:
        r_.cells[0].width = Inches(1.7); r_.cells[1].width = Inches(0.7); r_.cells[2].width = Inches(0.6)
        r_.cells[3].width = Inches(0.6); r_.cells[4].width = Inches(3.0); r_.cells[5].width = Inches(0.9)


VHEX = {"PASS": "1E7A32", "FAIL": "B31B1B", "UNVERIFIED": "8A8F98", "N/A": "6B7280", "PARTIAL": "B86A00"}
VCOL = {"PASS": GREEN, "FAIL": RED, "UNVERIFIED": GRAY, "N/A": LIGHT, "PARTIAL": AMBER}


def detail_section(doc):
    h2(doc, "Chi tiết test theo tiêu chí")
    intro = doc.add_paragraph()
    ir = intro.add_run("Mỗi tiêu chí = 1 kịch bản test cụ thể với kết quả mong đợi (expected), kết quả thực tế (actual) và ảnh chụp bằng chứng (proof). Verdict: ")
    ir.font.size = Pt(9.5); ir.font.color.rgb = GRAY
    for v, nm in [("PASS", "đạt"), ("FAIL", "lỗi thật"), ("UNVERIFIED", "chưa test được"), ("PARTIAL", "một phần"), ("N/A", "không áp dụng")]:
        rr = intro.add_run(f"{v}={nm}  "); rr.font.size = Pt(9); rr.bold = True; rr.font.color.rgb = VCOL[v]
    for c in ["C1", "C2", "C3", "C4", "C5", "C6", "C7"]:
        d = TD.get(c)
        if not d: continue
        verdict = d.get("verdict", "UNVERIFIED")
        ht = doc.add_table(rows=1, cols=1); hc = ht.rows[0].cells[0]
        shade(hc, VHEX.get(verdict, "8A8F98")); set_cell_margins(hc, 90, 90, 140, 140); hc.text = ""
        p = hc.paragraphs[0]; r = p.add_run(f"{d['tc']}  ·  {c} — {CRIT[c]}")
        r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = WHITE; r.font.name = 'Calibri'
        rv = p.add_run(f"    [{verdict}]"); rv.bold = True; rv.font.size = Pt(9)
        rv.font.color.rgb = RGBColor(0xFF, 0xE6, 0xE6) if verdict == "FAIL" else WHITE
        dt = doc.add_table(rows=3, cols=2); dt.style = 'Table Grid'
        labels = [("Kịch bản test", d.get("scenario", "")),
                  ("Kết quả mong đợi (expected)", d.get("expected", "")),
                  ("Kết quả thực tế (actual)", d.get("actual", ""))]
        for i, (lab, val) in enumerate(labels):
            lc = dt.rows[i].cells[0]; vc_ = dt.rows[i].cells[1]
            shade(lc, LIGHTBG_HEX); set_cell_margins(lc, 70, 70, 120, 120); set_cell_margins(vc_, 70, 70, 120, 120)
            cell_text(lc, lab, bold=True, color=NAVY, size=9)
            cell_text(vc_, val, size=9, color=(RED if (i == 2 and verdict == "FAIL") else DARK))
            lc.width = Inches(1.7); vc_.width = Inches(5.3)
        shots = d.get("shots", [])
        if shots:
            sp = doc.add_paragraph(); sr = sp.add_run("Bằng chứng:"); sr.bold = True
            sr.font.size = Pt(8.5); sr.font.color.rgb = GRAY
            sp.paragraph_format.space_after = Pt(2); sp.paragraph_format.space_before = Pt(2)
            for sh in shots[:2]:
                add_image(doc, os.path.join(B, sh), sh, 4.6)
        else:
            sp = doc.add_paragraph(); sr = sp.add_run("Bằng chứng: không có ảnh (tiêu chí dựa trên dữ liệu HTTP/spec).")
            sr.italic = True; sr.font.size = Pt(8.5); sr.font.color.rgb = LIGHT
            sp.paragraph_format.space_after = Pt(6)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    t = PRODUCT
    doc = Document(); set_base_style(doc)
    sec = doc.sections[0]; add_page_number_footer(sec, t['name'])
    levels = t["levels"]; no_cb = (not t["has_chatbot"])
    final, contrib, used_w, w = aggregate(levels, no_cb)
    cov = used_w
    unver = sum(1 for l in levels.values() if l == 0)
    conf = 0.55 if unver >= 4 else (0.75 if unver >= 2 else 0.92)
    if cov < 60: band = "⚠ ĐỘ PHỦ THẤP — phần chưa kiểm tra"
    elif final >= 85: band = "🟢 XUẤT SẮC — sản phẩm tốt"
    elif final >= 70: band = "🟡 KHÁ — cần cải thiện"
    elif final >= 50: band = "🟠 TRUNG BÌNH — vấn đề đáng kể"
    else: band = "🔴 YẾU — chặn luồng chính"
    build_cover(doc, t, final, cov, conf, band)
    h2(doc, "Tổng quan"); p = doc.add_paragraph(t["summary"])
    for r in p.runs: r.font.size = Pt(10.5)
    h2(doc, "Bảng điểm theo tiêu chí (Rubric)")
    rubric_table(doc, levels, no_cb, t)
    note = doc.add_paragraph()
    rn = note.add_run(f"Độ phủ {cov}% — {7 - unver}/7 tiêu chí đã kiểm tra. Tiêu chí chưa kiểm tra được (UNV) do thiếu tài khoản hoặc tính năng bị khóa sau đăng nhập — được loại khỏi tổng điểm, không xếp 'kém'.")
    rn.italic = True; rn.font.size = Pt(8.5); rn.font.color.rgb = GRAY
    detail_section(doc)
    # Feature inventory
    h2(doc, f"Bảng tính năng phát hiện ({len(INVENTORY)} tính năng)")
    it = doc.add_table(rows=1, cols=5); it.style = 'Table Grid'
    for i, htext in enumerate(["Tính năng", "Route", "Vai trò", "Loại", "Trạng thái"]):
        shade(it.rows[0].cells[i], TEAL_HEX); set_cell_margins(it.rows[0].cells[i])
        cell_text(it.rows[0].cells[i], htext, bold=True, color=WHITE, size=8.5)
    for x in INVENTORY[:12]:
        rc = it.add_row().cells
        for cc in rc: set_cell_margins(cc)
        st = x.get('status', '')
        cell_text(rc[0], (x.get('feature') or '')[:34], size=8.5, bold=True)
        cell_text(rc[1], x.get('route', ''), size=8, color=GRAY)
        cell_text(rc[2], x.get('role', ''), size=8)
        cell_text(rc[3], x.get('type', ''), size=8)
        cell_text(rc[4], st, size=8, color=(GREEN if st == 'rendered' else RED))
    for r_ in it.rows:
        r_.cells[0].width = Inches(2.3); r_.cells[1].width = Inches(1.7)
        r_.cells[2].width = Inches(1.1); r_.cells[3].width = Inches(0.8); r_.cells[4].width = Inches(0.9)
    # Screenshots
    h2(doc, "Bằng chứng hình ảnh")
    seen = set()
    for c in ["C2", "C4", "C1", "C5", "C3", "C6"]:
        sh = t["shot"].get(c)
        if sh and sh not in seen:
            seen.add(sh)
            cap = {"C2": "Xác thực đăng nhập (verify token + role-nav)",
                   "C4": "Phản hồi trợ lý AI",
                   "C1": "Bảng tính năng / dashboard",
                   "C5": "Tình trạng tải / console",
                   "C3": "Chức năng cốt lõi",
                   "C6": "UX / responsive"}.get(c, "Bằng chứng")
            add_image(doc, os.path.join(B, sh), f"{sh} — {cap}", 6.0)
    # method note
    h2(doc, "Phương pháp & độ tin cậy")
    for line in [
        f"{unver}/7 tiêu chí ở mức chưa kiểm tra được (UNVERIFIED) — chủ yếu do thiếu tài khoản hoặc tính năng bị khóa. Các tiêu chí này được loại khỏi tổng điểm.",
        "Mỗi điểm đều dựa trên kết quả kiểm thử thực tế (HTTP/token/ảnh chụp), không theo danh tiếng.",
        "Ảnh chụp đã nhúng trực tiếp vào báo cáo.",
        f"So với lần trước: {t['prev']}.",
    ]:
        bp = doc.add_paragraph(style='List Bullet'); bp.add_run(line).font.size = Pt(9.5)
    os.makedirs(RD, exist_ok=True)
    out = os.path.join(RD, f"QA-Report-{t['name'].replace(' ', '-')}.docx")
    doc.save(out); print(f"BUILT {os.path.getsize(out) // 1024}KB -> {out}")


if __name__ == "__main__":
    build()
